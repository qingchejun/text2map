"""
Text2Map Backend - FastAPI 应用主入口
提供文本转思维导图的核心API服务
"""

import os
import srt
import io
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import google.generativeai as genai
from dotenv import load_dotenv

# --- 初始化与配置 ---

# 加载 .env 文件中的环境变量
load_dotenv()

# 配置 Gemini API
try:
    genai.configure(api_key=os.environ.get("GOOGLE_API_KEY"))
except TypeError:
    print("错误：未找到 GOOGLE_API_KEY。请确保您的 .env 文件配置正确。")
    # 在实际应用中，这里应该优雅地退出或处理
    
app = FastAPI(
    title="Text2Map API",
    description="一个将长文本智能转换为思维导图的后端服务。",
    version="1.0.0"
)

# 添加CORS中间件
origins = [
    "http://localhost:3000",
    "https://text2map-frontend.onrender.com",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,  # 允许前端域名
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有HTTP方法
    allow_headers=["*"],  # 允许所有请求头
)

# --- Pydantic 数据模型 (用于API数据校验) ---

class TextInput(BaseModel):
    text: str

class MindmapResponse(BaseModel):
    mindmap_data: str

# --- 核心辅助函数 ---

def parse_srt_content(srt_string: str) -> str:
    """解析SRT字符串，只提取纯文本内容。"""
    subtitles = srt.parse(srt_string)
    full_text = "\n".join(sub.content for sub in subtitles)
    return full_text

def parse_file_content(filename: str, file_bytes: bytes) -> str:
    """
    根据文件扩展名解析文件内容
    """
    import os
    
    # 获取文件扩展名
    file_extension = os.path.splitext(filename)[1].lower()
    
    # 检查文件大小（限制为50MB）
    if len(file_bytes) > 50 * 1024 * 1024:
        raise ValueError("文件大小超过50MB限制")
    
    # 根据扩展名选择解析方法
    if file_extension == '.txt':
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        for encoding in encodings:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return file_bytes.decode('latin-1', errors='ignore').strip()
        
    elif file_extension == '.md':
        try:
            return file_bytes.decode('utf-8').strip()
        except UnicodeDecodeError:
            return parse_file_content(filename, file_bytes)  # 回退到txt解析
            
    elif file_extension == '.docx':
        try:
            from docx import Document
            doc_stream = io.BytesIO(file_bytes)
            doc = Document(doc_stream)
            
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            return '\n'.join(paragraphs)
        except ImportError:
            raise ValueError("python-docx库未安装，无法解析DOCX文件")
        except Exception as e:
            raise ValueError(f"DOCX文件解析失败: {str(e)}")
            
    elif file_extension == '.pdf':
        try:
            import PyPDF2
            pdf_stream = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text.strip())
                except Exception as e:
                    print(f"PDF第{page_num + 1}页解析失败: {str(e)}")
                    continue
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ValueError("PyPDF2库未安装，无法解析PDF文件")
        except Exception as e:
            raise ValueError(f"PDF文件解析失败: {str(e)}")
            
    elif file_extension == '.srt':
        try:
            content = file_bytes.decode('utf-8')
            subtitle_generator = srt.parse(content)
            subtitles = list(subtitle_generator)
            
            texts = []
            for subtitle in subtitles:
                text = subtitle.content.strip()
                if text:
                    texts.append(text)
            
            return '\n'.join(texts)
        except UnicodeDecodeError:
            try:
                content = file_bytes.decode('gbk')
                subtitle_generator = srt.parse(content)
                subtitles = list(subtitle_generator)
                
                texts = []
                for subtitle in subtitles:
                    text = subtitle.content.strip()
                    if text:
                        texts.append(text)
                
                return '\n'.join(texts)
            except Exception:
                raise ValueError("SRT文件编码不支持")
        except Exception as e:
            raise ValueError(f"SRT文件解析失败: {str(e)}")
    else:
        supported_formats = ['.txt', '.md', '.docx', '.pdf', '.srt']
        raise ValueError(f"不支持的文件格式: {file_extension}。支持格式: {', '.join(supported_formats)}")

def generate_mindmap_data(text_content: str) -> str | None:
    """
    调用 Google Gemini API 来处理文本并生成思维导图Markdown。
    如果成功，返回Markdown字符串；如果失败，返回None。
    """
    PROMPT_TEMPLATE = """
    你是一个顶级的知识架构师和信息分析专家。你的核心任务是将用户提供的复杂、可能结构混乱的原始文本，转换成一份极其详细、高度结构化、完全忠于原文信息的 Markdown 格式思维导图。

    **你必须严格遵循以下所有规则：**

    1.  **【无损原则】**: 你的首要目标是 **零信息损失**。必须捕捉并包含原文中所有的关键概念、论点、论据、数据、案例和细节。如果原文提到了一个具体的名字、数字或例子，你的导图中也必须体现出来。

    2.  **【结构保留原则】**: 尽可能地识别并保留原文的内在逻辑结构和层级关系。如果原文是按"总-分-总"或者"问题-分析-解决"的结构来写的，你的思维导图主干也应该反映出这种结构。

    3.  **【逐层深化】**:
        * 一级标题 (#): 应该是整个文档最核心、最顶层的主题。
        * 二级标题 (##): 应该是支撑核心主题的关键分支或主要部分。
        * 三级标题 (###): 应该是对二级分支的进一步展开或子论点。
        * 列表项 (-): 用于列举具体的细节、例子、数据或步骤。可以使用多级缩进列表来表示更深层次的从属关系。

    4.  **【精确提炼，而非泛泛总结】**: 你的输出应该是对原文信息的**结构化呈现**，而不是模糊的概括。
        * **错误示范**: "作者讨论了几个工具。"
        * **正确示范**: "- 工具示例: Evernote, Notion"

    现在，请基于以上所有规则，处理以下原始文本：
    """
    
    try:
        print(f"API Key 状态: {'已设置' if os.environ.get('GOOGLE_API_KEY') else '未设置'}")
        print("正在初始化 Gemini 模型...")
        model = genai.GenerativeModel('gemini-1.5-flash-latest')
        print("模型初始化成功")
        
        print("正在调用 Gemini API...")
        full_prompt = PROMPT_TEMPLATE + text_content
        print(f"输入文本长度: {len(text_content)} 字符")
        
        response = model.generate_content(full_prompt)
        print("成功获取API响应。")
        print(f"响应长度: {len(response.text)} 字符")
        return response.text
    except Exception as e:
        print(f"调用 API 时发生错误: {e}")
        print(f"错误类型: {type(e).__name__}")
        import traceback
        traceback.print_exc()
        return None

# --- API 路由 (Endpoints) ---

@app.get("/", summary="服务根路径，用于健康检查")
def read_root():
    """检查API服务是否正在运行。"""
    return {"message": "Text2Map Backend is running!"}

@app.post("/generate", response_model=MindmapResponse, summary="生成思维导图")
def create_mindmap(text_input: TextInput):
    """
    接收用户提交的文本，调用AI模型生成思维导图，并返回Markdown格式的结果。
    """
    if not text_input.text or text_input.text.isspace():
        raise HTTPException(status_code=400, detail="输入的文本不能为空。")
        
    result = generate_mindmap_data(text_input.text)
    
    if result is None:
        raise HTTPException(status_code=500, detail="AI服务处理失败，请稍后再试。")
        
    return MindmapResponse(mindmap_data=result)

@app.post("/generate-from-file", response_model=MindmapResponse, summary="从文件生成思维导图")
def create_mindmap_from_file(file: UploadFile = File(...)):
    """
    接收上传的文件，解析文件内容，调用AI模型生成思维导图。
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    
    try:
        # 读取文件内容
        file_content = file.file.read()
        
        # 解析文件内容
        try:
            extracted_text = parse_file_content(file.filename, file_content)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            raise HTTPException(status_code=500, detail="文件解析失败，请检查文件格式是否正确")
        
        # 验证提取的文本
        if not extracted_text or extracted_text.isspace():
            raise HTTPException(status_code=400, detail="文件内容为空或无法提取有效文本")
        
        # 调用AI处理
        result = generate_mindmap_data(extracted_text)
        
        if result is None:
            raise HTTPException(status_code=500, detail="AI服务处理失败，请稍后再试")
        
        return MindmapResponse(mindmap_data=result)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail="服务器内部错误，请稍后再试") 