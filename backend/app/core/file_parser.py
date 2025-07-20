"""
文件解析模块
支持多种文件格式的文本提取功能
"""

import io
import os
from typing import Optional
import logging

# 导入必要的库
try:
    from docx import Document
except ImportError:
    Document = None
    logging.warning("python-docx not installed, .docx files will not be supported")

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logging.warning("PyPDF2 not installed, .pdf files will not be supported")

try:
    import srt
except ImportError:
    srt = None
    logging.warning("srt not installed, .srt files will not be supported")


def parse_txt_content(file_bytes: bytes) -> str:
    """
    解析TXT文件内容
    
    Args:
        file_bytes: 文件字节内容
        
    Returns:
        提取的文本内容
    """
    try:
        # 尝试多种编码格式
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                content = file_bytes.decode(encoding)
                return content.strip()
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，使用latin-1作为最后手段
        return file_bytes.decode('latin-1', errors='ignore').strip()
        
    except Exception as e:
        raise ValueError(f"TXT文件解析失败: {str(e)}")


def parse_md_content(file_bytes: bytes) -> str:
    """
    解析Markdown文件内容
    
    Args:
        file_bytes: 文件字节内容
        
    Returns:
        提取的文本内容（保留Markdown格式）
    """
    try:
        # Markdown文件通常使用UTF-8编码
        content = file_bytes.decode('utf-8')
        return content.strip()
        
    except UnicodeDecodeError:
        # 如果UTF-8失败，尝试其他编码
        return parse_txt_content(file_bytes)
    except Exception as e:
        raise ValueError(f"Markdown文件解析失败: {str(e)}")


def parse_docx_content(file_bytes: bytes) -> str:
    """
    解析DOCX文件内容
    
    Args:
        file_bytes: 文件字节内容
        
    Returns:
        提取的文本内容
    """
    if Document is None:
        raise ValueError("python-docx库未安装，无法解析DOCX文件")
    
    try:
        # 创建内存文件对象
        doc_stream = io.BytesIO(file_bytes)
        doc = Document(doc_stream)
        
        # 提取所有段落的文本
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # 只添加非空段落
                paragraphs.append(text)
        
        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)
        
        return '\n'.join(paragraphs)
        
    except Exception as e:
        raise ValueError(f"DOCX文件解析失败: {str(e)}")


def parse_pdf_content(file_bytes: bytes) -> str:
    """
    解析PDF文件内容
    
    Args:
        file_bytes: 文件字节内容
        
    Returns:
        提取的文本内容
    """
    if PyPDF2 is None:
        raise ValueError("PyPDF2库未安装，无法解析PDF文件")
    
    try:
        # 创建内存文件对象
        pdf_stream = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        
        # 提取所有页面的文本
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            try:
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_parts.append(text.strip())
            except Exception as e:
                logging.warning(f"PDF第{page_num + 1}页解析失败: {str(e)}")
                continue
        
        return '\n\n'.join(text_parts)
        
    except Exception as e:
        raise ValueError(f"PDF文件解析失败: {str(e)}")


def parse_srt_content(file_bytes: bytes) -> str:
    """
    解析SRT字幕文件内容
    
    Args:
        file_bytes: 文件字节内容
        
    Returns:
        提取的文本内容（只包含字幕文本，不包含时间码）
    """
    if srt is None:
        raise ValueError("srt库未安装，无法解析SRT文件")
    
    try:
        # 解码文件内容
        content = file_bytes.decode('utf-8')
        
        # 解析SRT内容
        subtitle_generator = srt.parse(content)
        subtitles = list(subtitle_generator)
        
        # 提取所有字幕文本
        texts = []
        for subtitle in subtitles:
            text = subtitle.content.strip()
            if text:
                texts.append(text)
        
        return '\n'.join(texts)
        
    except UnicodeDecodeError:
        # 如果UTF-8失败，尝试其他编码
        try:
            content = file_bytes.decode('gbk')
            subtitle_generator = srt.parse(content)
            subtitles = list(subtitle_generator)
            
            texts = []
            for subtitle in subtitles:
                text = subtitle.content.strip()
                if text:
                    texts.append(text)
            
            return '\n'.join(texts)
        except Exception:
            raise ValueError("SRT文件编码不支持")
    except Exception as e:
        raise ValueError(f"SRT文件解析失败: {str(e)}")


def parse_file_content(filename: str, file_bytes: bytes) -> str:
    """
    根据文件扩展名自动选择解析函数
    
    Args:
        filename: 文件名（包含扩展名）
        file_bytes: 文件字节内容
        
    Returns:
        提取的文本内容
        
    Raises:
        ValueError: 当文件格式不支持或解析失败时
    """
    # 获取文件扩展名（小写）
    file_extension = os.path.splitext(filename)[1].lower()
    
    # 检查文件大小（限制为50MB）
    if len(file_bytes) > 50 * 1024 * 1024:
        raise ValueError("文件大小超过50MB限制")
    
    # 根据扩展名选择解析函数
    if file_extension == '.txt':
        return parse_txt_content(file_bytes)
    elif file_extension == '.md':
        return parse_md_content(file_bytes)
    elif file_extension == '.docx':
        return parse_docx_content(file_bytes)
    elif file_extension == '.pdf':
        return parse_pdf_content(file_bytes)
    elif file_extension == '.srt':
        return parse_srt_content(file_bytes)
    else:
        supported_formats = ['.txt', '.md', '.docx', '.pdf', '.srt']
        raise ValueError(f"不支持的文件格式: {file_extension}。支持格式: {', '.join(supported_formats)}")


def get_supported_formats() -> list:
    """
    获取支持的文件格式列表
    
    Returns:
        支持的文件格式列表
    """
    formats = ['.txt', '.md']
    
    if Document is not None:
        formats.append('.docx')
    
    if PyPDF2 is not None:
        formats.append('.pdf')
    
    if srt is not None:
        formats.append('.srt')
    
    return formats 